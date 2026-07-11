import hashlib
import re
import tempfile
from pathlib import Path
from typing import Any

from bs4 import BeautifulSoup

from .base import RawDocument


SUPPORTED_EXTENSIONS = {".csv", ".docx", ".epub", ".htm", ".html", ".md", ".pdf", ".txt"}


def stable_cloud_filename(file_id: str, name: str) -> str:
    suffix = Path(name).suffix.lower() or ".txt"
    digest = hashlib.sha256(file_id.encode("utf-8")).hexdigest()[:16]
    return f"cloud-{digest}{suffix}"


def _html_text(data: bytes) -> str:
    soup = BeautifulSoup(data, "html.parser")
    for element in soup(["script", "style"]):
        element.decompose()
    blocks = []
    for element in soup.find_all(["h1", "h2", "h3", "h4", "p", "li", "blockquote", "pre"]):
        text = element.get_text(" ", strip=True)
        if not text:
            continue
        if element.name and element.name.startswith("h"):
            blocks.append(f"{'#' * int(element.name[1])} {text}")
        elif element.name == "li":
            blocks.append(f"- {text}")
        else:
            blocks.append(text)
    return "\n\n".join(blocks)


def normalize_cloud_bytes(name: str, data: bytes) -> tuple[str, list[dict[str, Any]] | None]:
    extension = Path(name).suffix.casefold()
    if extension in {".md", ".txt", ".csv"}:
        return data.decode("utf-8", errors="replace").strip(), None
    if extension in {".htm", ".html"}:
        return _html_text(data), None
    if extension == ".pdf":
        import pymupdf

        segments = []
        with pymupdf.open(stream=data, filetype="pdf") as document:
            for index, page in enumerate(document):
                text = page.get_text("text").strip()
                if text:
                    segments.append(
                        {
                            "ordinal": index,
                            "locator": {"type": "pdf_page", "value": str(index + 1)},
                            "heading": f"Page {index + 1}",
                            "text": text,
                        }
                    )
        content = "\n\n".join(
            f"## {segment['heading']}\n\n{segment['text']}" for segment in segments
        )
        return content, segments
    if extension == ".epub":
        import ebooklib  # type: ignore[import-untyped]
        from ebooklib import epub  # type: ignore[import-untyped]

        with tempfile.NamedTemporaryFile(suffix=".epub") as file:
            file.write(data)
            file.flush()
            book = epub.read_epub(file.name, options={"ignore_ncx": True})
        segments = []
        for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
            text = _html_text(item.get_body_content())
            if text:
                segments.append(
                    {
                        "ordinal": len(segments),
                        "locator": {"type": "epub_href", "value": item.get_name()},
                        "heading": Path(item.get_name()).stem,
                        "text": text,
                    }
                )
        content = "\n\n".join(
            f"## {segment['heading']}\n\n{segment['text']}" for segment in segments
        )
        return content, segments
    if extension == ".docx":
        from docx import Document  # type: ignore[import-untyped]

        with tempfile.NamedTemporaryFile(suffix=".docx") as file:
            file.write(data)
            file.flush()
            document = Document(file.name)
        blocks = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style = paragraph.style.name.casefold() if paragraph.style else ""
            if style.startswith("heading"):
                level = re.search(r"\d+", style)
                blocks.append(f"{'#' * int(level.group() if level else '2')} {text}")
            else:
                blocks.append(text)
        return "\n\n".join(blocks), None
    raise ValueError(f"Unsupported cloud file extension '{extension or '(none)'}'.")


def make_cloud_document(
    *,
    kb_id: int,
    source_id: int,
    provider: str,
    file_id: str,
    name: str,
    data: bytes,
    modified_at: float,
    metadata: dict[str, Any],
) -> RawDocument:
    content, segments = normalize_cloud_bytes(name, data)
    if not content:
        raise ValueError(f"Cloud file '{name}' contains no indexable text.")
    raw_bytes = content.encode("utf-8")
    return RawDocument(
        kb_id=kb_id,
        source_id=source_id,
        source_type="cloud_drive",
        origin_path=f"{provider}:{file_id}",
        filename=stable_cloud_filename(file_id, name),
        title=Path(name).stem,
        format=Path(name).suffix.lstrip(".").lower() or "txt",
        raw_bytes=raw_bytes,
        content_markdown=content,
        content_hash=hashlib.sha256(raw_bytes).hexdigest(),
        source_modified_at=modified_at,
        metadata={"provider": provider, "provider_file_id": file_id, **metadata},
        segments=segments,
    )
